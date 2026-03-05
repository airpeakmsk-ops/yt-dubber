from __future__ import annotations
import pytest
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from yt_dubber.docx_exporter import write_review_docx, read_review_docx
from yt_dubber.models import SubtitleSegment, SegmentStatus


def make_segments():
    return [
        SubtitleSegment(index=1, start_sec=1.0, end_sec=3.0, duration_sec=2.0,
                        jp_text="テスト", ru_text="Тестовый текст",
                        status=SegmentStatus.TRANSLATED),
        SubtitleSegment(index=2, start_sec=4.0, end_sec=6.0, duration_sec=2.0,
                        jp_text="♪", ru_text="", status=SegmentStatus.SKIPPED),
        SubtitleSegment(index=3, start_sec=7.0, end_sec=9.0, duration_sec=2.0,
                        jp_text="魚", ru_text="Рыба", status=SegmentStatus.TRANSLATED),
    ]


def test_write_creates_file(tmp_path):
    path = str(tmp_path / "review.docx")
    segments = make_segments()
    result = write_review_docx(segments, path)
    assert result == path
    assert (tmp_path / "review.docx").exists()


def test_write_correct_column_count(tmp_path):
    path = str(tmp_path / "review.docx")
    write_review_docx(make_segments(), path)
    doc = DocxDocument(path)
    table = doc.tables[0]
    assert len(table.columns) == 4


def test_write_header_row(tmp_path):
    path = str(tmp_path / "review.docx")
    write_review_docx(make_segments(), path)
    doc = DocxDocument(path)
    table = doc.tables[0]
    hdr = table.rows[0].cells
    assert hdr[0].text == "#"
    assert hdr[1].text == "Timecode"
    assert hdr[2].text == "JP Original"
    assert "RU Translation" in hdr[3].text


def test_write_data_rows(tmp_path):
    path = str(tmp_path / "review.docx")
    segments = make_segments()
    write_review_docx(segments, path)
    doc = DocxDocument(path)
    table = doc.tables[0]
    # First data row (segment index=1)
    row1 = table.rows[1].cells
    assert row1[0].text == "1"
    assert "0:01" in row1[1].text  # start_sec=1.0 → 0:01
    assert "0:03" in row1[1].text  # end_sec=3.0 → 0:03
    assert row1[2].text == "テスト"
    assert row1[3].text == "Тестовый текст"


def test_write_grey_shading(tmp_path):
    path = str(tmp_path / "review.docx")
    write_review_docx(make_segments(), path)
    doc = DocxDocument(path)
    table = doc.tables[0]
    # Header row col 0 should have grey shading
    cell = table.cell(0, 0)
    tc = cell._tc
    tcPr = tc.tcPr
    shd_elements = tcPr.findall(qn("w:shd")) if tcPr is not None else []
    assert len(shd_elements) > 0, "Header col 0 should have shading"
    shd = shd_elements[0]
    assert shd.get(qn("w:fill")) == "D9D9D9", "Header col 0 should be grey D9D9D9"
    # RU column header (col 3) is also grey
    cell3 = table.cell(0, 3)
    tc3 = cell3._tc
    tcPr3 = tc3.tcPr
    shd_elements3 = tcPr3.findall(qn("w:shd")) if tcPr3 is not None else []
    # Data row col 3 (RU) should NOT be grey
    data_cell_ru = table.cell(1, 3)
    tc_ru = data_cell_ru._tc
    tcPr_ru = tc_ru.tcPr
    if tcPr_ru is not None:
        shd_ru = tcPr_ru.findall(qn("w:shd"))
        for shd in shd_ru:
            fill = shd.get(qn("w:fill"))
            assert fill != "D9D9D9", "Data row RU col should NOT have grey shading"


def test_write_skipped_segment_appears(tmp_path):
    path = str(tmp_path / "review.docx")
    segments = make_segments()
    write_review_docx(segments, path)
    doc = DocxDocument(path)
    table = doc.tables[0]
    # 3 segments → 3 data rows + 1 header = 4 rows total
    assert len(table.rows) == 4
    # Row 2 (index 1 in data rows) is the SKIPPED segment with empty ru_text
    row2 = table.rows[2].cells
    assert row2[0].text == "2"
    assert row2[3].text == ""


def test_roundtrip_basic(tmp_path):
    path = str(tmp_path / "review.docx")
    segments = make_segments()
    write_review_docx(segments, path)
    result = read_review_docx(path)
    assert 1 in result
    assert result[1] == "Тестовый текст"


def test_roundtrip_deleted_row(tmp_path):
    written_path = str(tmp_path / "original.docx")
    modified_path = str(tmp_path / "modified.docx")
    segments = make_segments()
    write_review_docx(segments, written_path)

    # Delete row at index 2 (segment with index=2)
    doc = DocxDocument(written_path)
    tbl = doc.tables[0]
    row_to_delete = tbl.rows[2]._tr  # rows[0]=header, rows[1]=seg1, rows[2]=seg2
    tbl._tbl.remove(row_to_delete)
    doc.save(modified_path)

    result = read_review_docx(modified_path)
    assert 1 in result, "Segment 1 should be present"
    assert 3 in result, "Segment 3 should be present"
    assert 2 not in result, "Segment 2 was deleted, should not be present"


def test_roundtrip_empty_ru(tmp_path):
    path = str(tmp_path / "review.docx")
    segments = make_segments()
    write_review_docx(segments, path)
    result = read_review_docx(path)
    # Segment 2 is SKIPPED with empty ru_text
    assert 2 in result
    assert result[2] == ""


def test_read_returns_int_keys(tmp_path):
    path = str(tmp_path / "review.docx")
    write_review_docx(make_segments(), path)
    result = read_review_docx(path)
    for key in result.keys():
        assert isinstance(key, int), f"Key {key!r} should be int, got {type(key)}"
