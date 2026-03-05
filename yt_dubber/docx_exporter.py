from __future__ import annotations
from typing import List, Dict
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from yt_dubber.models import SubtitleSegment

GREY_HEX = "D9D9D9"
HEADERS = [
    "#",
    "Timecode",
    "JP Original",
    "RU Translation\n(НЕ РЕДАКТИРОВАТЬ col 1-3 / DO NOT EDIT col 1-3)",
]
COL_WIDTHS = [Inches(0.4), Inches(1.1), Inches(2.25), Inches(2.25)]


def _shade_cell(cell, hex_color: str) -> None:
    """Apply background shading to a table cell using OxmlElement."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove any existing shading elements first
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    # Create a fresh OxmlElement per call (CRITICAL: reuse causes XML corruption)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")     # required
    shd.set(qn("w:color"), "auto")    # required
    shd.set(qn("w:fill"), hex_color)  # actual color
    tcPr.append(shd)


def _fmt(seconds: float) -> str:
    """Convert float seconds to 'M:SS' format for timecode display."""
    total = int(seconds)
    m, s = divmod(total, 60)
    return f"{m}:{s:02d}"


def _set_row_widths(cells, widths) -> None:
    """Set cell widths for a row."""
    for cell, width in zip(cells, widths):
        cell.width = width


def write_review_docx(segments: List[SubtitleSegment], output_path: str) -> str:
    """Write a 4-column review DOCX table: #, Timecode, JP Original, RU Translation.

    Columns 0-2 (index, timecode, JP) have grey shading (D9D9D9) to indicate read-only.
    Column 3 (RU translation) is white — the only editable column.
    Returns output_path.
    """
    doc = Document()
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.autofit = False

    # Set column-level widths
    for col, width in zip(table.columns, COL_WIDTHS):
        col.width = width
        for cell in col.cells:
            cell.width = width

    # Header row — grey all 4 cells
    hdr = table.rows[0].cells
    for i, (cell, header) in enumerate(zip(hdr, HEADERS)):
        cell.text = header
        cell.width = COL_WIDTHS[i]
        _shade_cell(cell, GREY_HEX)

    # Data rows
    for seg in segments:
        row = table.add_row()
        cells = row.cells
        timecode = f"{_fmt(seg.start_sec)}\u2013{_fmt(seg.end_sec)}"
        cells[0].text = str(seg.index)
        cells[1].text = timecode
        cells[2].text = seg.jp_text
        cells[3].text = seg.ru_text
        # Grey on cols 0, 1, 2 (read-only); col 3 (RU) left white — no shading call
        for col_idx in (0, 1, 2):
            _shade_cell(cells[col_idx], GREY_HEX)
        # Set cell widths for this row
        for col_idx, width in enumerate(COL_WIDTHS):
            cells[col_idx].width = width

    doc.save(output_path)
    return output_path


def read_review_docx(path: str) -> Dict[int, str]:
    """Read back user-edited Russian translations from the DOCX review table.

    Returns dict mapping segment index (int, from col 0) -> edited RU text (str).
    Keys are the index values from column 0, not row positions.
    Deleted rows cause that index to be absent from the returned dict.
    """
    doc = Document(path)
    table = doc.tables[0]
    result: Dict[int, str] = {}
    for row in table.rows[1:]:  # skip header row
        cells = row.cells
        idx_text = cells[0].text.strip()
        ru_text = cells[3].text.strip()
        if not idx_text:
            continue
        try:
            idx = int(idx_text)
        except ValueError:
            continue
        result[idx] = ru_text
    return result
